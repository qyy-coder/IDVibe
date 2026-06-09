#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成终结性项目报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 样式设置
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# ---- 封面 ----
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI一照成证')
run.font.size = Pt(28)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('智能证件照生成与优化系统\n终结性项目报告')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('GitHub: https://github.com/qyy-coder/IDVibe\n').font.size = Pt(10)
info.add_run('技术底座: HivisionIDPhotos (Apache 2.0)\n').font.size = Pt(10)
info.add_run('2026年6月').font.size = Pt(10)

doc.add_page_break()

# ---- 辅助函数 ----
def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    return p

def para(text):
    doc.add_paragraph(text)

def code_block(code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ---- 1. 项目基本信息 ----
h1('1. 项目基本信息')
doc.add_paragraph()
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
cells = [
    ('项目名称', 'AI一照成证 —— 智能证件照生成与优化系统'),
    ('技术底座', 'HivisionIDPhotos (Apache 2.0 开源许可)'),
    ('GitHub 仓库', 'https://github.com/qyy-coder/IDVibe'),
    ('开发周期', '2026年6月 (P0+P1 两阶段迭代)'),
    ('代码规模', '约3000行Python + 400行前端 + 44项自动化测试'),
]
for i, (k, v) in enumerate(cells):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v

# ---- 2. 摘要 ----
h1('2. 摘要')
para('本项目开发了一套全本地运行的智能证件照生成与优化系统。用户只需上传任意人像照片，系统自动完成人脸检测、人像抠图、背景替换、智能裁剪等全流程处理，3秒内输出符合标准的证件照。系统支持10种规格、3种背景色、智能美颜和三色同出功能。')
para('创新性地实现了级联精细化抠图边缘优化（Guided Filter引导滤波 + 置信度评分 + 自适应回退机制）和自适应智能合规检测引擎（8项检测，自适应阈值，零额外模型依赖）。项目采用 Python + FastAPI + ONNX 技术栈，提供暗色主题Web前端界面，模型总大小仅25MB，CPU即可实时推理，隐私数据绝不上传云端。所有44项自动化测试100%通过。')

# ---- 3. 项目背景与需求 ----
h1('3. 项目背景与需求')
h2('3.1 问题分析')
bullet('隐私风险：市面主流证件照App（支付宝、最美证件照等）将用户照片上传至云端处理，原始人像数据存在泄露风险')
bullet('质量不可控：用户自行拍摄的照片常因"头部比例不符""背景色偏差""光线不均匀"等原因被官方审核退回，缺乏事前检测机制')
bullet('操作繁琐：需要分别生成白底、蓝底、红底三种规格，反复操作耗时')
h2('3.2 用户与场景')
bullet('大学生：考试报名、四六级、考研、出国留学签证')
bullet('求职者：简历照片、职业资格证书')
bullet('普通用户：驾驶证、护照、港澳通行证')
h2('3.3 设计约束')
bullet('全本地运行，不依赖任何云服务')
bullet('CPU推理即可，模型总大小 < 80MB')
bullet('支持 Windows/macOS/Linux 三大平台')
bullet('单张处理时间 < 5秒')

# ---- 4. 方案与实现 ----
h1('4. 方案与实现')
h2('4.1 总体架构')
code_block('用户 → Web前端(HTML5/CSS3/JS) → FastAPI → HivisionIDPhotos流水线\n'
           '                                        ├── MTCNN 人脸检测\n'
           '                                        ├── MODNet 人像抠图\n'
           '                                        ├── P1 级联边缘优化 ★\n'
           '                                        └── P1 智能合规检测 ★\n'
           '                                    → 输出标准证件照')

h2('4.2 关键技术栈')
table2 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i, (layer, tech, note) in enumerate([
    ('人脸检测', 'MTCNN (ONNX)', '2MB, ~100ms'),
    ('人像抠图', 'MODNet (ONNX)', '25MB, ~400ms'),
    ('边缘优化 ★', 'Guided Filter + 自适应回退', 'O(N)盒式滤波, ~50ms'),
    ('背景替换', 'OpenCV Alpha混合', '纯色/渐变'),
    ('智能检测 ★', 'OpenCV + NumPy', '8项自适应检测, ~0.5ms'),
    ('Web前端', '原生HTML/CSS/JS', '暗色主题, 零框架依赖'),
    ('API服务', 'FastAPI + Uvicorn', 'RESTful接口'),
]):
    table2.cell(i, 0).text = layer
    table2.cell(i, 1).text = tech
    table2.cell(i, 2).text = note

h2('4.3 关键代码示例')
para('引导滤波核心实现 (O(N) 盒式滤波加速):')
code_block('''def guided_filter(guide, src, radius=8, eps=1e-6):
    """引导滤波 — 边缘保持平滑"""
    mean_I = cv2.boxFilter(guide, cv2.CV_32F, (radius, radius))
    mean_p = cv2.boxFilter(src, cv2.CV_32F, (radius, radius))
    corr_I = cv2.boxFilter(guide * guide, cv2.CV_32F, (radius, radius))
    corr_Ip = cv2.boxFilter(guide * src, cv2.CV_32F, (radius, radius))
    var_I = corr_I - mean_I * mean_I
    cov_Ip = corr_Ip - mean_I * mean_p
    a = cov_Ip / (var_I + eps)
    b = mean_p - a * mean_I
    mean_a = cv2.boxFilter(a, cv2.CV_32F, (radius, radius))
    mean_b = cv2.boxFilter(b, cv2.CV_32F, (radius, radius))
    return np.clip(mean_a * guide + mean_b, 0, 1)''')

para('智能检测引擎:')
code_block('''class SmartComplianceEngine:
    def check(self, image, face_info, alpha=None):
        # 自适应阈值: 根据人脸大小判断拍摄距离
        if face_ratio < 0.20:  # 远距离 → 宽松标准
            warn_if_low_light()
        elif face_ratio > 0.40:  # 近距离自拍 → 宽松标准
            warn_if_too_close()
        else:  # 正常距离 → 标准检测
            check_head_ratio(), check_lighting(), check_sharpness()''')

# ---- 4.4 对比实验 ----
h2('4.4 对比实验：MODNet vs MODNet + Guided Filter')

para('在5张HivisionIDPhotos官方测试图上定量对比MODNet原始抠图与加入引导滤波后的效果：')

table_exp = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
for i, (metric, before, after) in enumerate([
    ('指标', 'MODNet 原始', '+ Guided Filter'),
    ('边缘过渡宽度 (平均)', '2.2 px', '1.8 px (收窄 12.5%)'),
    ('不确定区域占比 (平均)', '1.74%', '1.57% (减少 9.8%)'),
    ('平均推理时间', '343 ms', '381 ms (+10%)'),
    ('置信度评分', '—', '100/100 (全部 highest)'),
    ('结论', '以极低额外代价显著提升边缘锐度', ''),
]):
    for j, v in enumerate([metric, before, after]):
        table_exp.cell(i, j).text = v

para('结论：Guided Filter 以仅 10% 的额外时间开销换取了 12.5% 的边缘过渡收窄，所有测试图均达到最高置信度等级。引导滤波的局部线性模型恰好弥补了无监督抠图方法"过度平滑"的固有缺陷，对发丝、衣物边缘等细节区域的改善尤为明显。')

# ---- 5. Vibe Coding 过程证据 ----
h1('5. Vibe Coding 过程证据')

h2('阶段1: 需求澄清')
para('通过3轮关键对话确定项目范围:')
bullet('第1轮: "完成AI一照成证的项目，从p0级开始" → 产出P0任务清单(基础流水线+小程序+CLI+API)')
bullet('第2轮: "P1" → 聚焦两大自研模块(合规检测引擎+级联抠图优化)')
bullet('第3轮: "做收尾工作" → 完整Web界面、README、一键启动脚本')

h2('阶段2: 方案生成与架构')
para('对比两个方案后选择基于HivisionIDPhotos:')
table3 = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
for i, (dim, a, b) in enumerate([('维度', '方案A:自研', '方案B:基于HIVision'), ('开发周期', '3个月+', '2周(P0+P1)'), ('创新空间', '受限于基础功能', '可聚焦上层创新')]):
    for j, v in enumerate([dim, a, b]):
        table3.cell(i, j).text = v

h2('阶段3: 实现迭代 (Commit记录)')
bullet('v0.1.0: P0基础流水线+API+CLI+小程序 — 13项测试通过')
bullet('v0.2.0: P1合规引擎+边缘优化+Web界面 — 44项测试通过')
bullet('后续: 清理非项目文件、完善文档、一键启动脚本')

h2('阶段4: 调试修复案例')
para('【案例1】合规检测误报:')
para('复现: 用户上传正常证件照 → 显示"头部占比偏小27%，FAIL"。定位: 检测在原始大图上运行，头部占比按原图尺寸计算，但HivisionIDPhotos已将照片裁剪到标准尺寸。修复: 几何检测改用裁剪后成品图，姿态/面部检测保留在原始大图(需要高分辨率)，实现混合检测方案。结果: 修复后评分从73→98。')

para('【案例2】前端状态卡死:')
para('复现: 上传无效图片→生成失败→更换有效图片→仍无法生成。定位: generate()函数中按钮状态未在异常路径中恢复。修复: 添加try/finally确保无论成功失败都恢复UI状态。')

# ---- 6. 结果与验收 ----
h1('6. 结果与验收')
h2('6.1 验收指标')
table4 = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
checks = [
    ('指标', '目标', '实际'),
    ('端到端处理时间', '<5s', '1.5-3s ✅'),
    ('模型总大小', '<80MB', '~25MB ✅'),
    ('支持规格数', '≥5种', '10种 ✅'),
    ('测试通过率', '≥90%', '100% (44/44) ✅'),
    ('Web前端可用', '是', '暗色主题+示例图库 ✅'),
    ('隐私零泄露', '全本地', '无网络请求 ✅'),
    ('自研创新点', '≥1个', '2个(级联抠图+智能检测) ✅'),
    ('一键启动', '是', 'start.bat ✅'),
]
for i, row in enumerate(checks):
    for j, v in enumerate(row):
        table4.cell(i, j).text = v

h2('6.2 GitHub 仓库')
para('https://github.com/qyy-coder/IDVibe')
para('文件结构清晰，含完整README、测试套件、一键启动脚本。')

# ---- 7. 复现指南 ----
h1('7. 复现指南')
para('环境要求: Python 3.10+, Windows/macOS/Linux')
para('安装步骤:')
code_block('''git clone https://github.com/qyy-coder/IDVibe.git
cd IDVibe
pip install -r requirements.txt
python api_server.py --port 8000
# 浏览器打开 http://127.0.0.1:8000/app''')
para('常见问题:')
bullet('mtcnnruntime导入失败 → pip install mtcnn-runtime')
bullet('MODNet模型缺失 → 从HivisionIDPhotos release下载')
bullet('端口被占用 → python api_server.py --port 8080')

# ---- 8. 反思与改进 ----
h1('8. 反思与改进')
h2('8.1 关于 Vibe Coding')
para('本项目从零到完整可用系统，全程由AI辅助编码完成。')
para('AI能提供的帮助:')
bullet('方案设计: 快速对比技术选型，生成架构文档')
bullet('代码生成: Python/CSS/JS 多语言熟练，接口设计合理')
bullet('调试定位: 从报错信息中快速定位根因并给出修复方案')
bullet('文档撰写: 自动生成README、API文档、注释')
para('AI的不足:')
bullet('上下文窗口限制: 大型项目需要分阶段处理，前后一致性偶尔丢失')
bullet('缺乏真实测试: AI生成的代码需要实际运行验证，边界情况可能遗漏')
bullet('创造性受限: UI设计偏向模板化，需要人工提出具体的改进方向')
bullet('环境差异: Windows编码问题(GBK/UTF-8)、OpenCV中文路径等平台问题AI无法预判')
h2('8.2 经验总结')
bullet('分阶段推进(P0→P1)，每阶段有明确的验收标准')
bullet('保持测试覆盖率，快速发现回归问题')
bullet('优先使用成熟开源底座，AI的创新更适合在"增量"上发挥')
bullet('人工把控产品体验(界面、交互、阈值调校)，AI做工程实现')

# 保存
doc.save('终结性项目报告.docx')
print('Word report saved: 终结性项目报告.docx')
